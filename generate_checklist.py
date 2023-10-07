from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Initialize a new Word Document
doc = Document()
doc.add_heading('Data Validation Checklist: January 2021 to December 2022', 0)

# Tasks for data validation with categories and checkboxes
tasks_with_categories_checkboxes = {
    "File-Level Validation": [
        "Confirm file existence and accessibility.",
        "Validate file format.",
        "Check file completeness for the month."
    ],
    "Header Validation": [
        "Validate column names and order.",
        "Rename `member_casual` to `user_type`."
    ],
    "Data Type and Format Validation": [
        "Ensure data type consistency for each column.",
        "Validate and convert date formats in started_at and ended_at.",
        "Validate format of start_station_id and end_station_id."
    ],
    "Content Validation": [
        "Validate values in rideable_type.",
        "Ensure user_type contains only \"member\" or \"casual\".",
        "Check correspondence between station names and IDs.",
        "Validate uniqueness of ride_id.",
        "Ensure valid geographical coordinates."
    ],
    "Consistency and Logical Checks": [
        "Verify logical consistency in ride durations.",
        "Check spatial consistency between start and end locations."
    ],
    "Missing Value Check": [
        "Identify and handle missing values."
    ],
    "Outlier Detection": [
        "Detect and handle outliers in ride durations and spatial data."
    ],
    "Station ID Range": [
        "Validate start_station_id and end_station_id against sensible ranges or a master list."
    ],
    "Temporal Consistency": [
        "Identify suspicious temporal patterns.",
        "Ensure no future dates are present in the data."
    ],
    "Documentation": [
        "Log all steps, inconsistencies, and actions taken during validation."
    ]
}

# Define the years and months for the checklist
years = [2021, 2022]
months = [
    "January", "February", "March", "April", "May", "June",
    "July", "August", "September", "October", "November", "December"
]

# Define the checkbox character
checkbox_char = "[ ] "

# Loop through each month and year, adding the tasks to the document
for year in years:
    for month in months:
        doc.add_heading(f'{month} {year} Data Validation Checklist', level=1)
        for i, (category, tasks) in enumerate(tasks_with_categories_checkboxes.items(), start=1):
            doc.add_heading(f"{i}. {category}:", level=2)
            for j, task in enumerate(tasks, start=1):
                paragraph = doc.add_paragraph(style='BodyText')
                run = paragraph.add_run(f"{checkbox_char}{j}. {task}")
                run.bold = True
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# Save the document
file_path_checklist_formatted = 'data_validation_checklist_formatted.docx'
doc.save(file_path_checklist_formatted)

# If you want to print the path
print(file_path_checklist_formatted)
